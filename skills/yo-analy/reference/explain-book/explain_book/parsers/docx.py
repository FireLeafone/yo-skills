from __future__ import annotations

import zipfile
import sys
from explain_book.exceptions import ExtractionError


def extract_docx_with_python_docx(docx_path: str) -> str | None:
    # 本函数被无条件调用（而非仅经由 extract_docx()），因此在已安装
    # python-docx 的情况下被直接调用时也能自我保护：在 python-docx 打开
    # 压缩包之前，就因 DOCTYPE/ENTITY 声明抛出 ExtractionError。如果未安装
    # python-docx，则完全不做校验、直接返回 None —— 未安装的解析器什么也
    # 解析不了，跳过扫描不会损失任何安全性（无论恶意与否都提取不到内容），
    # 也避免在（默认的、仅标准库的）情况下为每次 extract_docx() 调用都支付
    # 一次完整压缩包扫描 —— 那种情况下本解析器根本不会运行。直接调用本
    # 函数、且无论 python-docx 是否可用都需要校验保证的调用方，应改用
    # extract_docx_with_zipfile()，或自行调用 validate_docx_xml_safety()。
    try:
        import docx
        validate_docx_xml_safety(docx_path)
        document = docx.Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        return None
    except ExtractionError:
        # 没有这一句，下面宽泛的 `except Exception` 会把
        # validate_docx_xml_safety() 的 XXE 拒绝也一并捕获，把一次安全拒绝
        # 变成被吞掉的 [warn] + None。
        raise
    except Exception as e:
        print(f"  [warn] extract_docx_with_python_docx failed: {type(e).__name__}: {e}", file=sys.stderr)
        return None


def extract_docx_with_zipfile(docx_path: str) -> str | None:
    # 本函数被无条件调用（而非仅经由 extract_docx()），因此被直接调用时也
    # 能自我保护：在 XML 到达解析器之前，就因 DOCTYPE/ENTITY 声明抛出
    # ExtractionError。
    validate_docx_xml_safety(docx_path)
    try:
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(docx_path) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts: list[str] = []

        def emit_block(elem) -> None:
            # 按文档顺序遍历块级内容。段落合并其内部的 run；表格每行输出一行
            # 制表符连接的文本（行格式与 python-docx 路径相同，但保持原有顺序
            # —— python-docx 会把所有表格追加到最后）。未知的包裹元素（如
            # <w:sdt> 内容控件）会递归进入，以免其中的段落/表格丢失；<w:p> 和
            # <w:tbl> 则不递归，避免表格单元格里的段落被重复计数。单元格文本
            # 由单元格内的 run 拼接而成；嵌套表格归入父单元格，同时也会单独
            # 输出（少见；尽力而为）。
            for child in elem:
                tag = child.tag
                if tag == f"{ns}p":
                    texts = [t.text for t in child.iter(f"{ns}t") if t.text]
                    if texts:
                        parts.append("".join(texts))
                elif tag == f"{ns}tbl":
                    for row in child.iter(f"{ns}tr"):
                        cells = []
                        for cell in row.iter(f"{ns}tc"):
                            cell_texts = [t.text for t in cell.iter(f"{ns}t") if t.text]
                            cells.append("".join(cell_texts).strip())
                        if any(cells):
                            parts.append("\t".join(cells))
                else:
                    emit_block(child)

        body = root.find(f"{ns}body")
        emit_block(body if body is not None else root)
        return "\n".join(parts) if parts else None
    except Exception as e:
        print(f"  [warn] extract_docx_with_zipfile failed: {type(e).__name__}: {e}", file=sys.stderr)
        return None


def validate_docx_xml_safety(docx_path: str) -> None:
    """扫描 DOCX zip 压缩包中的所有 XML 文件，防止 XML 实体扩展（Billion Laughs）和 XXE 注入。"""
    try:
        with zipfile.ZipFile(docx_path) as zf:
            for name in zf.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    xml_bytes = zf.read(name)
                    for encoding in ("utf-8", "utf-16", "utf-16le", "utf-16be", "utf-32"):
                        try:
                            content = xml_bytes.decode(encoding, errors="ignore").upper()
                        except LookupError:
                            continue
                        if "<!DOCTYPE" in content or "<!ENTITY" in content:
                            raise ExtractionError(
                                f"Security validation failed: XML file '{name}' in DOCX archive contains forbidden DTD or entity declarations."
                            )
    except zipfile.BadZipFile as e:
        raise ExtractionError(f"Invalid DOCX file: {e}")
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Error during security validation of DOCX archive: {e}")


def extract_docx(docx_path: str) -> tuple[str, str]:
    # 校验放在各个叶子解析器（extract_docx_with_python_docx、
    # extract_docx_with_zipfile）中，这样无论实际由哪个解析器处理文件，
    # 校验都恰好执行一次，而不是在这里执行一次、再在最终落入的解析器里
    # 又执行一次。
    print("Trying python-docx...", end=" ", flush=True)
    text = extract_docx_with_python_docx(docx_path)
    if text and text.strip():
        print("OK")
        return text, "python-docx"

    print("not available")
    print("Trying stdlib DOCX parser...", end=" ", flush=True)
    text = extract_docx_with_zipfile(docx_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile-docx"

    print("FAILED")
    raise ExtractionError(
        "Could not extract text from DOCX.\n"
        "Install python-docx for best results:\n"
        "  pip3 install python-docx"
    )
